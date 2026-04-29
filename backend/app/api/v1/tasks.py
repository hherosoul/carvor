import json
import logging
from pathlib import Path
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Query
from fastapi.responses import PlainTextResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, delete
from pydantic import BaseModel

from app.core.database import get_session
from app.models.database import Task, TaskReference, Paper, Idea, IdeaReference, Conversation, ConversationMessage, Experiment
from app.pipelines.experiment import upload_experiment_log, analyze_experiment
from app.pipelines.prompt_doc import save_prompt_doc
from app.gateway.llm_gateway import gateway
from app.core.config import DATA_DIR

router = APIRouter(prefix="/api/tasks", tags=["tasks"])


class TaskCreate(BaseModel):
    name: str
    source_idea_id: int


class RefAdd(BaseModel):
    paper_id: int


@router.get("")
async def list_tasks(session: AsyncSession = Depends(get_session)):
    result = await session.execute(select(Task).order_by(Task.created_at.desc()))
    tasks = result.scalars().all()
    return [{"id": t.id, "name": t.name, "research_goal": t.research_goal, "source_idea_id": t.source_idea_id, "created_at": t.created_at} for t in tasks]


@router.post("")
async def create_task(data: TaskCreate, session: AsyncSession = Depends(get_session)):
    result = await session.execute(select(Idea).where(Idea.id == data.source_idea_id))
    idea = result.scalar_one_or_none()
    if not idea:
        raise HTTPException(404, "Source idea not found")

    task = Task(
        name=data.name,
        research_goal=idea.content,
        source_idea_id=idea.id,
    )
    session.add(task)
    await session.flush()

    task_dir = DATA_DIR / "tasks" / f"task-{task.id}"
    task_dir.mkdir(parents=True, exist_ok=True)

    if idea.content:
        (task_dir / "idea.md").write_text(idea.content, encoding="utf-8")

    idea_refs = await session.execute(
        select(IdeaReference).where(IdeaReference.idea_id == idea.id)
    )
    for ref in idea_refs.scalars().all():
        task_ref = TaskReference(task_id=task.id, paper_id=ref.paper_id)
        session.add(task_ref)

    idea.status = "已立项"
    session.add(idea)
    await session.commit()

    return {"id": task.id, "name": task.name, "research_goal": task.research_goal}


@router.get("/{task_id}")
async def get_task(task_id: int, session: AsyncSession = Depends(get_session)):
    result = await session.execute(select(Task).where(Task.id == task_id))
    task = result.scalar_one_or_none()
    if not task:
        raise HTTPException(404, "Task not found")
    idea_content = None
    if task.source_idea_id:
        idea_result = await session.execute(select(Idea).where(Idea.id == task.source_idea_id))
        idea = idea_result.scalar_one_or_none()
        if idea:
            idea_content = idea.content
    return {"id": task.id, "name": task.name, "research_goal": task.research_goal, "source_idea_id": task.source_idea_id, "idea_content": idea_content, "created_at": task.created_at}


@router.delete("/{task_id}")
async def delete_task(task_id: int, session: AsyncSession = Depends(get_session)):
    result = await session.execute(select(Task).where(Task.id == task_id))
    task = result.scalar_one_or_none()
    if not task:
        raise HTTPException(404, "Task not found")
    await session.execute(delete(TaskReference).where(TaskReference.task_id == task_id))
    await session.delete(task)
    await session.commit()
    return {"ok": True}


@router.get("/{task_id}/references")
async def get_references(task_id: int, session: AsyncSession = Depends(get_session)):
    result = await session.execute(
        select(TaskReference, Paper)
        .join(Paper, Paper.id == TaskReference.paper_id)
        .where(TaskReference.task_id == task_id)
    )
    rows = result.all()
    return [{
        "paper_id": p.id,
        "title": p.title,
        "authors": json.loads(p.authors) if p.authors else [],
        "bibtex": r.bibtex,
        "tags": r.tags,
    } for r, p in rows]


@router.post("/{task_id}/references")
async def add_reference(task_id: int, data: RefAdd, session: AsyncSession = Depends(get_session)):
    existing = await session.execute(
        select(TaskReference).where(
            TaskReference.task_id == task_id,
            TaskReference.paper_id == data.paper_id,
        )
    )
    if existing.scalar_one_or_none():
        return {"ok": True, "message": "Already in reference pool"}
    ref = TaskReference(task_id=task_id, paper_id=data.paper_id)
    session.add(ref)
    await session.commit()
    return {"ok": True}


@router.delete("/{task_id}/references/{paper_id}")
async def remove_reference(task_id: int, paper_id: int, session: AsyncSession = Depends(get_session)):
    result = await session.execute(
        select(TaskReference).where(
            TaskReference.task_id == task_id,
            TaskReference.paper_id == paper_id,
        )
    )
    ref = result.scalar_one_or_none()
    if not ref:
        raise HTTPException(404, "Reference not found")
    await session.delete(ref)
    await session.commit()
    return {"ok": True}


@router.post("/{task_id}/references/{paper_id}/bibtex")
async def generate_bibtex(task_id: int, paper_id: int, session: AsyncSession = Depends(get_session)):
    result = await session.execute(select(Paper).where(Paper.id == paper_id))
    paper = result.scalar_one_or_none()
    if not paper:
        raise HTTPException(404, "Paper not found")

    llm_result = await gateway.call_async("bibtex_generate", {
        "title": paper.title,
        "authors": paper.authors or "",
    })

    bibtex = llm_result.get("bibtex", "")

    ref_result = await session.execute(
        select(TaskReference).where(
            TaskReference.task_id == task_id,
            TaskReference.paper_id == paper_id,
        )
    )
    ref = ref_result.scalar_one_or_none()
    if ref:
        ref.bibtex = bibtex
        session.add(ref)
        await session.commit()

    return {"bibtex": bibtex}


@router.post("/{task_id}/references/{paper_id}/tags")
async def recommend_tags(task_id: int, paper_id: int, session: AsyncSession = Depends(get_session)):
    result = await session.execute(select(Paper).where(Paper.id == paper_id))
    paper = result.scalar_one_or_none()
    if not paper:
        raise HTTPException(404, "Paper not found")

    task_result = await session.execute(select(Task).where(Task.id == task_id))
    task = task_result.scalar_one_or_none()

    llm_result = await gateway.call_async("tag_recommend", {
        "paper_title": paper.title,
        "paper_abstract": paper.abstract or "",
        "research_direction": task.research_goal if task else "",
    })

    tags = llm_result.get("tags", [])
    tags_str = json.dumps(tags, ensure_ascii=False)

    ref_result = await session.execute(
        select(TaskReference).where(
            TaskReference.task_id == task_id,
            TaskReference.paper_id == paper_id,
        )
    )
    ref = ref_result.scalar_one_or_none()
    if ref:
        ref.tags = tags_str
        session.add(ref)
        await session.commit()

    return {"tags": tags}


@router.get("/{task_id}/research")
async def get_research_doc(task_id: int):
    path = DATA_DIR / "tasks" / f"task-{task_id}" / "research.md"
    if path.exists():
        return {"content": path.read_text(encoding="utf-8")}
    return {"content": ""}


class SaveDocRequest(BaseModel):
    content: str


@router.put("/{task_id}/research")
async def save_research_doc(task_id: int, data: SaveDocRequest):
    logger = logging.getLogger("carvor.tasks")
    path = DATA_DIR / "tasks" / f"task-{task_id}" / "research.md"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(data.content, encoding="utf-8")
    logger.info(f"Research doc saved for task {task_id}, {len(data.content)} chars")
    return {"ok": True}


@router.get("/{task_id}/review")
async def get_review_doc(task_id: int):
    path = DATA_DIR / "tasks" / f"task-{task_id}" / "review.md"
    if path.exists():
        return {"content": path.read_text(encoding="utf-8")}
    return {"content": ""}


@router.put("/{task_id}/review")
async def save_review_doc(task_id: int, data: SaveDocRequest):
    path = DATA_DIR / "tasks" / f"task-{task_id}" / "review.md"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(data.content, encoding="utf-8")
    return {"ok": True}


@router.get("/{task_id}/method")
async def get_method_doc(task_id: int):
    path = DATA_DIR / "tasks" / f"task-{task_id}" / "method.md"
    if path.exists():
        return {"content": path.read_text(encoding="utf-8")}
    return {"content": ""}


@router.put("/{task_id}/method")
async def save_method_doc(task_id: int, data: SaveDocRequest):
    path = DATA_DIR / "tasks" / f"task-{task_id}" / "method.md"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(data.content, encoding="utf-8")
    return {"ok": True}


@router.get("/{task_id}/polish")
async def get_polish_doc(task_id: int):
    path = DATA_DIR / "tasks" / f"task-{task_id}" / "polish.md"
    if path.exists():
        return {"content": path.read_text(encoding="utf-8")}
    return {"content": ""}


@router.put("/{task_id}/polish")
async def save_polish_doc(task_id: int, data: SaveDocRequest):
    path = DATA_DIR / "tasks" / f"task-{task_id}" / "polish.md"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(data.content, encoding="utf-8")
    return {"ok": True}


@router.post("/{task_id}/polish/upload-docx")
async def upload_polish_docx(task_id: int, file: UploadFile = File(...)):
    if not file.filename:
        raise HTTPException(400, "No filename provided")
    ext = Path(file.filename).suffix.lower()
    if ext not in (".docx", ".doc"):
        raise HTTPException(400, "Only .docx and .doc files are supported")

    content = await file.read()

    if ext == ".doc":
        try:
            content = content.decode("utf-8", errors="replace")
        except Exception:
            pass

    try:
        from docx import Document
        import io
        doc = Document(io.BytesIO(content))
    except Exception as e:
        raise HTTPException(400, f"Failed to parse docx: {str(e)}")

    markdown_parts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            markdown_parts.append("")
            continue
        style_name = (para.style.name or "").lower() if para.style else ""
        if "heading 1" in style_name:
            markdown_parts.append(f"# {text}")
        elif "heading 2" in style_name:
            markdown_parts.append(f"## {text}")
        elif "heading 3" in style_name:
            markdown_parts.append(f"### {text}")
        elif "heading 4" in style_name:
            markdown_parts.append(f"#### {text}")
        elif "title" in style_name:
            markdown_parts.append(f"# {text}")
        else:
            markdown_parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            markdown_parts.append("| " + " | ".join(cells) + " |")

    markdown_content = "\n\n".join(markdown_parts)

    path = DATA_DIR / "tasks" / f"task-{task_id}" / "polish.md"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(markdown_content, encoding="utf-8")

    return {"content": markdown_content, "filename": file.filename}


@router.post("/{task_id}/generate-research")
async def generate_research_doc(task_id: int):
    task_dir = DATA_DIR / "tasks" / f"task-{task_id}"
    review_path = task_dir / "review.md"
    method_path = task_dir / "method.md"

    review_content = review_path.read_text(encoding="utf-8") if review_path.exists() else ""
    method_content = method_path.read_text(encoding="utf-8") if method_path.exists() else ""

    if not review_content and not method_content:
        return {"ok": False, "message": "综述和方法文档均为空，请先完成综述讨论和方法讨论"}

    parts = []
    if review_content:
        parts.append("# 文献综述\n\n" + review_content)
    if method_content:
        parts.append("# 研究方法\n\n" + method_content)

    research_content = "\n\n---\n\n".join(parts)
    research_path = task_dir / "research.md"
    research_path.write_text(research_content, encoding="utf-8")

    return {"ok": True, "content": research_content}


@router.get("/{task_id}/conversations")
async def list_conversations(task_id: int, session: AsyncSession = Depends(get_session)):
    result = await session.execute(
        select(Conversation).where(Conversation.task_id == task_id).order_by(Conversation.updated_at.desc())
    )
    convs = result.scalars().all()
    return [{"id": c.id, "scenario": c.scenario, "created_at": c.created_at} for c in convs]


@router.get("/{task_id}/experiments")
async def list_experiments(task_id: int, session: AsyncSession = Depends(get_session)):
    result = await session.execute(
        select(Experiment).where(Experiment.task_id == task_id).order_by(Experiment.created_at.desc())
    )
    exps = result.scalars().all()
    return [{"id": e.id, "log_path": e.log_path, "filename": Path(e.log_path).name if e.log_path else "", "analysis_report": e.analysis_report, "created_at": e.created_at} for e in exps]


@router.post("/{task_id}/experiments")
async def upload_experiment(
    task_id: int,
    file: UploadFile = File(...),
    session: AsyncSession = Depends(get_session),
):
    content = await file.read()
    text_content = None
    for enc in ("utf-8", "gbk", "gb2312", "latin-1"):
        try:
            text_content = content.decode(enc)
            break
        except (UnicodeDecodeError, LookupError):
            continue
    if text_content is None:
        text_content = content.decode("utf-8", errors="replace")
    experiment = await upload_experiment_log(task_id, text_content, file.filename or "exp.json", session)
    return {"id": experiment.id, "log_path": experiment.log_path}


@router.post("/{task_id}/experiments/{exp_id}/analyze")
async def analyze_exp(exp_id: int, session: AsyncSession = Depends(get_session)):
    report = await analyze_experiment(exp_id, session)
    return {"report": report}


@router.get("/{task_id}/experiments/{exp_id}")
async def get_experiment(exp_id: int, session: AsyncSession = Depends(get_session)):
    result = await session.execute(select(Experiment).where(Experiment.id == exp_id))
    exp = result.scalar_one_or_none()
    if not exp:
        raise HTTPException(404, "Experiment not found")
    log_path = Path(exp.log_path)
    log_content = log_path.read_text(encoding="utf-8") if log_path.exists() else ""
    return {"id": exp.id, "log_content": log_content, "analysis_report": exp.analysis_report, "filename": log_path.name}


@router.delete("/{task_id}/experiments/{exp_id}")
async def delete_experiment(exp_id: int, session: AsyncSession = Depends(get_session)):
    result = await session.execute(select(Experiment).where(Experiment.id == exp_id))
    exp = result.scalar_one_or_none()
    if not exp:
        raise HTTPException(404, "Experiment not found")
    log_path = Path(exp.log_path)
    if log_path.exists():
        log_path.unlink()
    await session.delete(exp)
    await session.commit()
    return {"ok": True}


@router.get("/{task_id}/prompt-docs")
async def list_prompt_docs(task_id: int):
    doc_dir = DATA_DIR / "tasks" / f"task-{task_id}" / "prompt-docs"
    if not doc_dir.exists():
        return []
    return [{"filename": f.name, "content": f.read_text(encoding="utf-8")} for f in doc_dir.glob("*.md")]


_DOC_MAP = {
    "research": "research.md",
    "review": "review.md",
    "method": "method.md",
    "polish": "polish.md",
}


@router.get("/{task_id}/export/{doc_type}")
async def export_doc(task_id: int, doc_type: str):
    if doc_type not in _DOC_MAP:
        raise HTTPException(400, f"Unknown doc type: {doc_type}. Available: {list(_DOC_MAP.keys())}")
    path = DATA_DIR / "tasks" / f"task-{task_id}" / _DOC_MAP[doc_type]
    if not path.exists():
        raise HTTPException(404, f"Document not found: {doc_type}")
    content = path.read_text(encoding="utf-8")
    filename = f"{doc_type}_{task_id}.md"
    return PlainTextResponse(
        content,
        media_type="text/markdown; charset=utf-8",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/{task_id}/export-prompt-doc/{filename}")
async def export_prompt_doc(task_id: int, filename: str):
    path = DATA_DIR / "tasks" / f"task-{task_id}" / "prompt-docs" / filename
    if not path.exists():
        raise HTTPException(404, "Prompt doc not found")
    content = path.read_text(encoding="utf-8")
    return PlainTextResponse(
        content,
        media_type="text/markdown; charset=utf-8",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
